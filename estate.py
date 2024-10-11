from PyQt5.QtGui import *
from PyQt5.QtWidgets import *
from PyQt5.QtCore import *
from PyQt5 import QtCore
from PyQt5.uic import loadUiType
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_TAB_ALIGNMENT,WD_PARAGRAPH_ALIGNMENT
import sys
import datetime
from datetime import date
import mysql.connector
import win32api
import arabic_reshaper
import bidi.algorithm

real_estateUI,_ = loadUiType('real_estate.ui')

class Main(QMainWindow, real_estateUI):
    def __init__(self,parent=None):
        super(Main,self).__init__(parent)
        QMainWindow.__init__(self)
        self.setupUi(self)
        
        self.tabWidget.tabBar().setVisible(False)
        self.dateEdit.setDate(QDate.currentDate())  
        self.dateEdit_2.setDate(QDate.currentDate())
        self.dateEdit_10.setDate(QDate.currentDate())
        self.dateEdit_11.setDate(QDate.currentDate())
        self.dateEdit_12.setDate(QDate.currentDate())
        self.dateEdit_6.setDate(QDate.currentDate())        
        self.pushButton_7.setEnabled(False)
        self.pushButton_10.setEnabled(False)
        self.pushButton_13.setEnabled(False)
        self.pushButton_20.setEnabled(False)
        self.pushButton_15.setEnabled(False)
        self.pushButton_11.setEnabled(False)
        self.pushButton_16.setEnabled(False)  
        self.pushButton_17.setEnabled(False)      
        self.pushButton_21.setEnabled(False)
        self.pushButton_22.setEnabled(False)  
        self.pushButton_23.setEnabled(False)
        self.pushButton_24.setEnabled(False)
        self.pushButton_25.setEnabled(False)
        self.pushButton_26.setEnabled(False)
        self.pushButton_29.setEnabled(False)
        self.checkBox.setVisible(False)

        self.lineEdit_37.setValidator(QIntValidator())
        self.lineEdit_37.setMaxLength(4)
        self.lineEdit_40.setValidator(QIntValidator())
        self.lineEdit_40.setMaxLength(4)

        self.groupBox.setEnabled(False)

        self.tabWidget.setCurrentIndex(4)


        self.db_connect()
        self.handel_buttons()
        self.aroud_search()
        self.egarat_search()
        self.talabat_search()
        self.tamleek_search()

    def db_connect(self):        
        self.db = mysql.connector.connect(user='root', password='',
                host='localhost', db='estates_db')
        self.cur = self.db.cursor(buffered=True)

        sql = f''' SELECT * from login'''
        self.cur.execute(sql)
        data = self.cur.fetchall()
        if data == []:
            self.checkBox.setVisible(True)
            self.pushButton_29.setEnabled(True)


    def handel_buttons(self):
        
        self.setTabOrder(self.lineEdit_17, self.lineEdit_37)
        self.setTabOrder(self.lineEdit, self.lineEdit_2)
        self.setTabOrder(self.lineEdit_2, self.lineEdit_3)
        self.setTabOrder(self.lineEdit_3, self.lineEdit_4)
        self.setTabOrder(self.lineEdit_4, self.lineEdit_5)
        self.setTabOrder(self.lineEdit_5, self.lineEdit_7)
        self.setTabOrder(self.lineEdit_7, self.lineEdit_8)
        self.setTabOrder(self.lineEdit_8, self.lineEdit_6)
        self.setTabOrder(self.lineEdit_14, self.lineEdit_15)
        self.setTabOrder(self.lineEdit_15, self.lineEdit_16)
        self.setTabOrder(self.lineEdit_16, self.lineEdit_10)
        self.setTabOrder(self.lineEdit_10, self.lineEdit_12)
        self.setTabOrder(self.lineEdit_12, self.lineEdit_11)
        self.setTabOrder(self.lineEdit_11, self.lineEdit_17)
        
        self.setTabOrder(self.lineEdit_73, self.lineEdit_68)
        self.setTabOrder(self.lineEdit_68, self.lineEdit_69)
        self.setTabOrder(self.lineEdit_69, self.lineEdit_72)
        self.setTabOrder(self.lineEdit_72, self.lineEdit_75)
        self.setTabOrder(self.lineEdit_75, self.lineEdit_76)
        self.setTabOrder(self.lineEdit_76, self.lineEdit_70)
        self.setTabOrder(self.lineEdit_70, self.lineEdit_74)
        self.setTabOrder(self.lineEdit_74, self.lineEdit_71)
        self.setTabOrder(self.lineEdit_71, self.lineEdit_67)

        self.setTabOrder(self.lineEdit_28, self.lineEdit_30)
        self.setTabOrder(self.lineEdit_30, self.lineEdit_26)
        self.setTabOrder(self.lineEdit_26, self.lineEdit_27)
        self.setTabOrder(self.lineEdit_27, self.lineEdit_31)
        self.setTabOrder(self.lineEdit_31, self.lineEdit_29)
        self.setTabOrder(self.lineEdit_29, self.lineEdit_33)
        self.setTabOrder(self.lineEdit_33, self.lineEdit_32)
        self.setTabOrder(self.lineEdit_32, self.lineEdit_48)
        self.setTabOrder(self.lineEdit_48, self.lineEdit_49)
        self.setTabOrder(self.lineEdit_49, self.lineEdit_35)        
        self.setTabOrder(self.lineEdit_35, self.lineEdit_36)
        self.setTabOrder(self.lineEdit_36, self.lineEdit_34)
        
        self.pushButton.clicked.connect(self.aroud_tap)
        self.pushButton_2.clicked.connect(self.talabat_tap)
        self.pushButton_3.clicked.connect(self.egarat_tap)
        self.pushButton_4.clicked.connect(self.tamleek_tap)
        self.pushButton_5.clicked.connect(self.aroud_save)
        self.pushButton_6.clicked.connect(self.talabat_save)
        self.pushButton_13.clicked.connect(self.talabat_clrscreen)
        self.pushButton_7.clicked.connect(self.talabat_edit)
        self.pushButton_18.clicked.connect(self.egarat_save)
        self.pushButton_20.clicked.connect(self.egarat_clrscreen)
        self.pushButton_8.clicked.connect(self.tamleek_save)
        self.pushButton_15.clicked.connect(self.tamleek_clrscreen)
        self.pushButton_9.clicked.connect(self.aroud_filer)
        self.pushButton_10.clicked.connect(self.aroud_clrscreen)
        self.pushButton_13.clicked.connect(self.talabat_clrscreen)
        self.pushButton_20.clicked.connect(self.egarat_clrscreen)
        self.pushButton_15.clicked.connect(self.tamleek_clrscreen)
        self.pushButton_11.clicked.connect(self.aroud_edit)
        self.pushButton_12.clicked.connect(self.talabat_filter)
        self.pushButton_21.clicked.connect(self.contract_print)
        self.pushButton_19.clicked.connect(self.egarat_filter)
        self.pushButton_22.clicked.connect(self.egarat_edit)
        self.pushButton_14.clicked.connect(self.tamleek_filter)
        self.pushButton_16.clicked.connect(self.tamleek_edit)
        self.pushButton_17.clicked.connect(self.tamleek_print)
        self.pushButton_23.clicked.connect(self.aroud_return)
        self.pushButton_24.clicked.connect(self.talabat_return)
        self.pushButton_25.clicked.connect(self.egarat_return)
        self.pushButton_26.clicked.connect(self.tamleek_return)
        self.pushButton_27.clicked.connect(self.login)
        self.pushButton_29.clicked.connect(self.signup)
        self.pushButton_28.clicked.connect(self.signout)

        self.lineEdit_38.textChanged.connect(self.aroud_srch_by_name)
        self.lineEdit_39.textChanged.connect(self.talabat_srch_by_name)
        self.lineEdit_78.textChanged.connect(self.egarat_srch_by_name)
        self.lineEdit_41.textChanged.connect(self.tamleek_srch_by_name)

        self.lineEdit_45.textChanged.connect(self.aroud_srch_by_category)
        self.lineEdit_46.textChanged.connect(self.talabat_srch_by_category)
        self.lineEdit_79.textChanged.connect(self.egarat_srch_by_category)
        self.lineEdit_47.textChanged.connect(self.tamleek_srch_by_category)

        self.lineEdit.textEdited.connect(self.enabled_aroud_edit)
        self.lineEdit_2.textEdited.connect(self.enabled_aroud_edit)
        self.lineEdit_3.textEdited.connect(self.enabled_aroud_edit)
        self.lineEdit_4.textEdited.connect(self.enabled_aroud_edit)
        self.lineEdit_5.textEdited.connect(self.enabled_aroud_edit)
        self.lineEdit_6.textEdited.connect(self.enabled_aroud_edit)
        self.lineEdit_7.textEdited.connect(self.enabled_aroud_edit)
        self.lineEdit_8.textEdited.connect(self.enabled_aroud_edit)

        self.lineEdit_10.textEdited.connect(self.enabled_talabat_edit)
        self.lineEdit_11.textEdited.connect(self.enabled_talabat_edit)
        self.lineEdit_12.textEdited.connect(self.enabled_talabat_edit)
        self.lineEdit_14.textEdited.connect(self.enabled_talabat_edit)
        self.lineEdit_15.textEdited.connect(self.enabled_talabat_edit)
        self.lineEdit_16.textEdited.connect(self.enabled_talabat_edit)
        self.lineEdit_17.textEdited.connect(self.enabled_talabat_edit)

        self.lineEdit_73.textEdited.connect(self.enabled_egarat_edit)
        self.lineEdit_68.textEdited.connect(self.enabled_egarat_edit)
        self.lineEdit_69.textEdited.connect(self.enabled_egarat_edit)
        self.lineEdit_72.textEdited.connect(self.enabled_egarat_edit)
        self.lineEdit_75.textEdited.connect(self.enabled_egarat_edit)
        self.lineEdit_76.textEdited.connect(self.enabled_egarat_edit)
        self.lineEdit_70.textEdited.connect(self.enabled_egarat_edit)
        self.lineEdit_74.textEdited.connect(self.enabled_egarat_edit)
        self.lineEdit_71.textEdited.connect(self.enabled_egarat_edit)
        self.lineEdit_67.textEdited.connect(self.enabled_egarat_edit)

        self.lineEdit_28.textEdited.connect(self.enabled_tamleek_edit)
        self.lineEdit_30.textEdited.connect(self.enabled_tamleek_edit)
        self.lineEdit_26.textEdited.connect(self.enabled_tamleek_edit)
        self.lineEdit_27.textEdited.connect(self.enabled_tamleek_edit)
        self.lineEdit_31.textEdited.connect(self.enabled_tamleek_edit)
        self.lineEdit_29.textEdited.connect(self.enabled_tamleek_edit)
        self.lineEdit_32.textEdited.connect(self.enabled_tamleek_edit)
        self.lineEdit_33.textEdited.connect(self.enabled_tamleek_edit)
        self.lineEdit_34.textEdited.connect(self.enabled_tamleek_edit)
        self.lineEdit_35.textEdited.connect(self.enabled_tamleek_edit)
        self.lineEdit_36.textEdited.connect(self.enabled_tamleek_edit)
        self.lineEdit_48.textEdited.connect(self.enabled_tamleek_edit)
        self.lineEdit_49.textEdited.connect(self.enabled_tamleek_edit)
        self.comboBox_2.currentTextChanged.connect(self.enabled_tamleek_edit)
    
    def aroud_srch_by_name(self):        
        self.lineEdit_37.setText('')

    def talabat_srch_by_name(self):
        self.lineEdit_40.setText('')

    def egarat_srch_by_name(self):
        self.lineEdit_77.setText('')

    def tamleek_srch_by_name(self):
        self.lineEdit_42.setText('')

    def aroud_srch_by_category(self):        
        self.lineEdit_37.setText('')
        self.lineEdit_38.setText('')

    def talabat_srch_by_category(self):
        self.lineEdit_40.setText('')
        self.lineEdit_39.setText('')

    def egarat_srch_by_category(self):
        self.lineEdit_77.setText('')
        self.lineEdit_78.setText('')

    def tamleek_srch_by_category(self):
        self.lineEdit_42.setText('')
        self.lineEdit_41.setText('')

    def enabled_aroud_edit(self):
        if self.pushButton_5.isEnabled() == False:            
            self.pushButton_11.setEnabled(True)        

    def enabled_talabat_edit(self):
        if self.pushButton_6.isEnabled() == False:        
            self.pushButton_7.setEnabled(True)

    def enabled_egarat_edit(self):
        if self.pushButton_18.isEnabled() == False:        
            self.pushButton_22.setEnabled(True)

    def enabled_tamleek_edit(self):
        if self.pushButton_8.isEnabled() == False:
            self.pushButton_16.setEnabled(True)

    def aroud_tap(self):
        self.pushButton.setStyleSheet("background-color : #c2fc03")
        self.pushButton_2.setStyleSheet("background-color : ")
        self.pushButton_3.setStyleSheet("background-color : ")
        self.pushButton_4.setStyleSheet("background-color : ")

        self.tabWidget.setCurrentIndex(0)

    def aroud_clrscreen(self):

        self.lineEdit.setText('')
        self.lineEdit_2.setText('')
        self.lineEdit_3.setText('')
        self.lineEdit_4.setText('')
        self.lineEdit_5.setText('')
        self.lineEdit_6.setText('')
        self.lineEdit_7.setText('')
        self.lineEdit_8.setText('')
        self.pushButton_5.setEnabled(True)
        self.pushButton_10.setEnabled(False)
        self.pushButton_11.setEnabled(False)

    def aroud_save(self):
        
        ar_name = self.lineEdit.text()
        ar_phone = self.lineEdit_2.text()
        ar_date = self.dateEdit.date()
        ar_date = ar_date.toString(QtCore.Qt.ISODate)        
        ar_category = self.lineEdit_3.text()
        ar_contract = self.lineEdit_4.text()
        ar_region = self.lineEdit_5.text()
        ar_address = self.lineEdit_8.text()
        ar_amount_required = self.lineEdit_7.text()
        ar_discription = self.lineEdit_6.text()

        if ar_name =='' or ar_phone =='' or ar_category =='' or ar_contract =='' or ar_region =='' or ar_address =='' or ar_amount_required == '' or ar_discription =='':
            QMessageBox.warning(self, 'بيانات ناقصة', 'من فضلك تأكد من إدخال جميع البيانات', QMessageBox.Ok)
            return

        self.cur.execute('''
            INSERT INTO aroud(name, phone, date, category, contract, region, address, amount_required, discription)
            VALUES(%s, %s, %s, %s, %s, %s, %s, %s, %s)
            ''',(ar_name, ar_phone, ar_date, ar_category, ar_contract, ar_region, ar_address, ar_amount_required, ar_discription))

        self.db.commit()
        self.pushButton_5.setEnabled(False)
        self.pushButton_10.setEnabled(True)
        self.aroud_search()
        #self.tabWidget_2.setCurrentIndex(1)
        QMessageBox.warning(self, 'حفظ بيانات', 'تم حفظ البيانات بنجاح', QMessageBox.Ok)
        return

    def aroud_search(self):
        self.tableWidget.setRowCount(0)
        self.tableWidget.insertRow(0)
        self.cur.execute('''
        SELECT id, name, phone, date, category, contract, region FROM aroud
        ''')
        data = self.cur.fetchall()

        for row, form in enumerate(data):
            for col, item in enumerate(form):
                self.tableWidget.setItem(row, col, QTableWidgetItem(str(item)))
                col += 1
            row_pos = self.tableWidget.rowCount()
            self.tableWidget.insertRow(row_pos)

        self.pushButton_23.setEnabled(False)


    def aroud_filer(self):
        if self.lineEdit_37.text()=='' and self.lineEdit_38.text()=='' and self.lineEdit_45.text()=='':
            QMessageBox.warning(self, 'لا توجد بيانات',  'لا توجد بيانات يمكن البحث عنها', QMessageBox.Ok)
            return
        if self.lineEdit_37.text() != '':            
            ar_id = self.lineEdit_37.text()
            sql = f''' SELECT * FROM aroud WHERE id='{ar_id}' '''
                        
            self.cur.execute(sql)
            data = self.cur.fetchone()
            if data == None :
                QMessageBox.warning(self, 'لا توجد بيانات',  'لا توجد بيانات تخص المعلومات التي أدخلتها', QMessageBox.Ok)
                return

            self.tabWidget_2.setCurrentIndex(0)        
            self.lineEdit.setText(data[1])
            self.lineEdit_2.setText(str(data[2]))
            self.dateEdit.setDate(data[3])
            self.lineEdit_3.setText(data[4])        
            self.lineEdit_4.setText(data[5])
            self.lineEdit_5.setText(data[6])
            self.lineEdit_7.setText(data[8])
            self.lineEdit_8.setText(data[7])
            self.lineEdit_6.setText(data[9]) 

        else:            
            if self.lineEdit_38.text() != '':                
                ar_name = self.lineEdit_38.text()
                sql = f''' SELECT * FROM aroud WHERE name LIKE '%{ar_name}%' '''
            elif self.lineEdit_45.text() != '':
                category = self.lineEdit_45.text()
                sql = f''' SELECT * FROM aroud WHERE category LIKE '%{category}%' '''
            self.cur.execute(sql)
            data = self.cur.fetchall()            
            if data == [] :
                self.pushButton_23.setEnabled(True)
                QMessageBox.warning(self, 'لا توجد بيانات',  'لا توجد بيانات تخص المعلومات التي أدخلتها', QMessageBox.Ok)
                return
            
            self.tableWidget.setRowCount(0)
            self.tableWidget.insertRow(0)
            for row, form in enumerate(data):
                for col, item in enumerate(form):
                    self.tableWidget.setItem(row, col, QTableWidgetItem(str(item)))
                    col += 1
                row_pos = self.tableWidget.rowCount()
                self.tableWidget.insertRow(row_pos)        
        
        self.lineEdit_37.setText('')
        self.lineEdit_38.setText('')
        self.lineEdit_45.setText('')            
        self.pushButton_23.setEnabled(True)
        self.pushButton_5.setEnabled(False)
        self.pushButton_10.setEnabled(True)
    
    def aroud_edit(self):
        ar_id = self.lineEdit_37.text()
        ar_name = self.lineEdit.text()
        ar_phone = self.lineEdit_2.text()
        ar_date = self.dateEdit.date()
        ar_date = ar_date.toString(QtCore.Qt.ISODate)        
        ar_category = self.lineEdit_3.text()
        ar_contract = self.lineEdit_4.text()
        ar_region = self.lineEdit_5.text()
        ar_address = self.lineEdit_8.text()
        ar_amount_required = self.lineEdit_7.text()
        ar_discription = self.lineEdit_6.text()        
        sql = f''' UPDATE aroud SET name='{ar_name}',\
            phone='{ar_phone}', date='{ar_date}',\
            category='{ar_category}', contract='{ar_contract}',\
            region='{ar_region}', address='{ar_address}',\
            amount_required='{ar_amount_required}',\
            discription='{ar_discription}'\
            WHERE id='{ar_id}' '''

        self.cur.execute(sql)
        self.db.commit()

        self.lineEdit_37.setText('')
        self.aroud_search()
        QMessageBox.warning(self, 'تعديل بيانات', 'تم تعديل البيانات بنجاح', QMessageBox.Ok)
        return
    
    def aroud_return(self):        
        self.aroud_search()
        self.lineEdit_37.setText('')
        self.lineEdit_38.setText('')
        self.lineEdit_45.setText('')
        self.pushButton_23.setEnabled(False)
        self.lineEdit_37.setEnabled(True)


    def talabat_tap(self):

        self.pushButton.setStyleSheet("background-color : ")
        self.pushButton_2.setStyleSheet("background-color : #c2fc03")
        self.pushButton_3.setStyleSheet("background-color : ")
        self.pushButton_4.setStyleSheet("background-color : ")
        self.tabWidget.setCurrentIndex(1)
    
    def talabat_save(self):

        ta_name = self.lineEdit_14.text()
        ta_phone = self.lineEdit_15.text()
        ta_date = self.dateEdit_2.date()
        ta_date = ta_date.toString(QtCore.Qt.ISODate)
        ta_category = self.lineEdit_16.text()
        ta_contract = self.lineEdit_10.text()
        ta_region = self.lineEdit_12.text()
        ta_rental_limit = self.lineEdit_11.text()
        ta_discription = self.lineEdit_17.text()

        if ta_name =='' or ta_phone =='' or ta_category =='' or ta_contract =='' or ta_region =='' or ta_rental_limit =='' or ta_discription == '' :
            QMessageBox.warning(self, 'بيانات ناقصة', 'من فضلك تأكد من إدخال جميع البيانات', QMessageBox.Ok)
            return

        self.cur.execute('''
            INSERT INTO talabat(name, phone, date, category, contract, region, rental_limit, discription)
            VALUES(%s, %s, %s, %s, %s, %s, %s, %s)
            ''',(ta_name, ta_phone, ta_date, ta_category, ta_contract, ta_region, ta_rental_limit, ta_discription))

        self.db.commit()
        self.talabat_search()
        self.pushButton_6.setEnabled(False)
        self.pushButton_13.setEnabled(True)
        QMessageBox.warning(self, 'حفظ بيانات', 'تم حفظ البيانات بنجاح', QMessageBox.Ok)
        return

    def talabat_clrscreen(self):

        self.lineEdit_14.setText('')
        self.lineEdit_15.setText('')
        self.lineEdit_16.setText('')
        self.lineEdit_10.setText('')
        self.lineEdit_12.setText('')
        self.lineEdit_11.setText('')
        self.lineEdit_17.setText('')
        self.pushButton_6.setEnabled(True)
        self.pushButton_7.setEnabled(False)
        self.pushButton_13.setEnabled(False)

    def talabat_search(self):

        self.tableWidget_2.setRowCount(0)
        self.tableWidget_2.insertRow(0)
        self.cur.execute('''
        SELECT id, name, phone, date, category, contract, region, rental_limit FROM talabat
        ''')
        data = self.cur.fetchall()

        for row, form in enumerate(data):
            for col, item in enumerate(form):
                self.tableWidget_2.setItem(row, col, QTableWidgetItem(str(item)))
                col += 1
            row_pos = self.tableWidget_2.rowCount()
            self.tableWidget_2.insertRow(row_pos)

        self.pushButton_24.setEnabled(False)

    def talabat_filter(self):
        if self.lineEdit_40.text() == '' and self.lineEdit_39.text()=='' and self.lineEdit_46.text()=='':
            QMessageBox.warning(self, 'لا توجد بيانات',  'لا توجد بيانات يمكن البحث عنها', QMessageBox.Ok)
            return
        if self.lineEdit_40.text() != '':            
            ta_id = self.lineEdit_40.text()            
            sql =f''' SELECT id, name, phone, date, category, contract, region,  rental_limit, discription FROM talabat WHERE id={ta_id} '''
                        
            self.cur.execute(sql)
            data = self.cur.fetchone()
            if data == None :
                QMessageBox.warning(self, 'لا توجد بيانات',  'لا توجد بيانات تخص المعلومات التي أدخلتها', QMessageBox.Ok)
                return

            self.tabWidget_4.setCurrentIndex(0)        
            self.lineEdit_14.setText(data[1])
            self.lineEdit_15.setText(str(data[2]))
            self.dateEdit_2.setDate(data[3])
            self.lineEdit_16.setText(data[4])        
            self.lineEdit_10.setText(data[5])
            self.lineEdit_12.setText(data[6])
            self.lineEdit_11.setText(data[7])
            self.lineEdit_17.setText(data[8])

        else:            
            if self.lineEdit_39.text() != '':
                ta_name = self.lineEdit_39.text()
                sql = f''' SELECT * FROM talabat WHERE name LIKE '%{ta_name}%' '''
            elif self.lineEdit_46.text() != '':
                category = self.lineEdit_46.text()
                sql = f''' SELECT * FROM talabat WHERE category LIKE '%{category}%' '''
            self.cur.execute(sql)
            data = self.cur.fetchall()            
            if data == [] :
                self.pushButton_24.setEnabled(True)
                QMessageBox.warning(self, 'لا توجد بيانات',  'لا توجد بيانات تخص المعلومات التي أدخلتها', QMessageBox.Ok)
                return
            
            self.tableWidget_2.setRowCount(0)
            self.tableWidget_2.insertRow(0)
            for row, form in enumerate(data):
                for col, item in enumerate(form):
                    self.tableWidget_2.setItem(row, col, QTableWidgetItem(str(item)))
                    col += 1
                row_pos = self.tableWidget_2.rowCount()
                self.tableWidget_2.insertRow(row_pos)        
        self.lineEdit_39.setText('')
        self.lineEdit_40.setText('')
        self.lineEdit_46.setText('')            
        self.pushButton_24.setEnabled(True)
        self.pushButton_6.setEnabled(False)
        self.pushButton_13.setEnabled(True)        
        
    
    def talabat_edit(self):

        ta_id = self.lineEdit_40.text()
        ta_name = self.lineEdit_14.text()
        ta_phone = self.lineEdit_15.text()
        ta_date = self.dateEdit_2.date()
        ta_date = ta_date.toString(QtCore.Qt.ISODate)        
        ta_category = self.lineEdit_16.text()
        ta_contract = self.lineEdit_10.text()
        ta_region = self.lineEdit_12.text()        
        ta_rental_limit = self.lineEdit_11.text()
        ta_discription = self.lineEdit_17.text()        
        sql = f''' UPDATE talabat SET name='{ta_name}',\
            phone='{ta_phone}', date='{ta_date}',\
            category='{ta_category}', contract='{ta_contract}',\
            region='{ta_region}', rental_limit='{ta_rental_limit}',\
            discription='{ta_discription}'\
            WHERE id='{ta_id}' '''

        self.cur.execute(sql)
        self.db.commit()

        self.lineEdit_40.setText('')
        self.talabat_search()        
        QMessageBox.warning(self, 'تعديل بيانات', 'تم تعديل البيانات بنجاح', QMessageBox.Ok)
        return

    def talabat_return(self):        
        self.talabat_search()
        self.pushButton_24.setEnabled(False)
        self.lineEdit_40.setEnabled(True)
        self.lineEdit_39.setText('')
        self.lineEdit_40.setText('')
        self.lineEdit_46.setText('')



    def egarat_tap(self):

        self.pushButton.setStyleSheet("background-color : ")
        self.pushButton_2.setStyleSheet("background-color : ")
        self.pushButton_3.setStyleSheet("background-color : #c2fc03")
        self.pushButton_4.setStyleSheet("background-color : ")
        self.tabWidget.setCurrentIndex(2)

    def egarat_save(self):
        eg_owner = self.lineEdit_73.text()
        eg_owner_id = self.lineEdit_68.text()
        eg_tenant = self.lineEdit_69.text()        
        eg_tenant_id = self.lineEdit_72.text()
        eg_date = self.dateEdit_10.date()
        eg_date = eg_date.toString(QtCore.Qt.ISODate)
        eg_category = self.comboBox_4.currentText()
        eg_discription = self.lineEdit_75.text()
        eg_address = self.lineEdit_76.text()
        eg_cont_term = self.lineEdit_70.text()
        eg_start = self.dateEdit_12.date()
        eg_start = eg_start.toString(QtCore.Qt.ISODate)
        eg_end = self.dateEdit_11.date()
        eg_end = eg_end.toString(QtCore.Qt.ISODate)
        eg_rental_value = self.lineEdit_74.text()
        eg_insurance = self.lineEdit_71.text()
        eg_purpose = self.lineEdit_67.text()

        if eg_owner =='' or eg_owner_id =='' \
            or eg_tenant =='' or eg_tenant_id =='' \
            or eg_category =='' or eg_discription =='' \
            or eg_address == '' or eg_cont_term == '' \
            or eg_insurance == '' or eg_purpose == '' :
            QMessageBox.warning(self, 'بيانات ناقصة', 'من فضلك تأكد من إدخال جميع البيانات', QMessageBox.Ok)
            return

        self.cur.execute('''
            INSERT INTO rents(owner, owner_id, tenant, tenant_id, date, category, discription, address, cont_term, start, end, rental_value, insurance, purpose)
            VALUES(%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
            ''',(eg_owner, eg_owner_id, eg_tenant, eg_tenant_id, eg_date, eg_category, eg_discription, eg_address, eg_cont_term, eg_start, eg_end, eg_rental_value, eg_insurance, eg_purpose))

        self.db.commit()
        self.egarat_search()
        self.pushButton_18.setEnabled(False)
        self.pushButton_20.setEnabled(True)
        self.pushButton_21.setEnabled(True)

        QMessageBox.warning(self, 'حفظ بيانات', 'تم حفظ البيانات بنجاح', QMessageBox.Ok)
        return
        

    def egarat_clrscreen(self):

        self.lineEdit_73.setText('')
        self.lineEdit_68.setText('')
        self.lineEdit_69.setText('')
        self.lineEdit_72.setText('')
        self.lineEdit_75.setText('')
        self.lineEdit_76.setText('')
        self.lineEdit_70.setText('')
        self.lineEdit_74.setText('')
        self.lineEdit_71.setText('')
        self.lineEdit_67.setText('')
        self.pushButton_18.setEnabled(True)
        self.pushButton_20.setEnabled(False)
        self.pushButton_21.setEnabled(False)
        self.pushButton_22.setEnabled(False)

    def egarat_search(self):
        self.tableWidget_3.setRowCount(0)
        self.tableWidget_3.insertRow(0)
        self.cur.execute('''
        SELECT id, owner, tenant, date, category, cont_term, rental_value FROM rents
        ''')
        data = self.cur.fetchall()

        for row, form in enumerate(data):
            for col, item in enumerate(form):
                self.tableWidget_3.setItem(row, col, QTableWidgetItem(str(item)))
                col += 1
            row_pos = self.tableWidget_3.rowCount()
            self.tableWidget_3.insertRow(row_pos)

    def egarat_filter(self):
        if self.lineEdit_77.text() == '' and self.lineEdit_78.text()=='' and self.lineEdit_79.text()=='':
            QMessageBox.warning(self, 'لا توجد بيانات',  'لا توجد بيانات يمكن البحث عنها', QMessageBox.Ok)
            return
        if self.lineEdit_77.text() != '':            
            eg_id = self.lineEdit_77.text()
            sql = f''' SELECT * FROM rents WHERE id='{eg_id}' '''
                        
            self.cur.execute(sql)
            data = self.cur.fetchone()
            if data == None :
                QMessageBox.warning(self, 'لا توجد بيانات',  'لا توجد بيانات تخص المعلومات التي أدخلتها', QMessageBox.Ok)
                return

            self.tabWidget_3.setCurrentIndex(0)        
            self.lineEdit_73.setText(data[1])
            self.lineEdit_68.setText(str(data[2]))
            self.lineEdit_69.setText(data[3])
            self.lineEdit_72.setText(str(data[4]))
            self.dateEdit_10.setDate(data[5])                
            self.comboBox_4.setCurrentText(data[6])
            self.lineEdit_75.setText(data[7])
            self.lineEdit_76.setText(data[8])
            self.lineEdit_70.setText(data[9])
            self.dateEdit_12.setDate(data[10])
            self.dateEdit_11.setDate(data[11])
            self.lineEdit_74.setText(data[12])
            self.lineEdit_71.setText(data[13])
            self.lineEdit_67.setText(data[14])

        else:            
            if self.lineEdit_78.text() != '':
                eg_name = self.lineEdit_78.text()
                sql = f''' SELECT * FROM rents WHERE owner LIKE '%{eg_name}%' '''
            elif self.lineEdit_79.text() != '':
                category = self.lineEdit_79.text()
                sql = f''' SELECT * FROM rents WHERE category LIKE '%{category}%' '''
            self.cur.execute(sql)
            data = self.cur.fetchall()            
            if data == [] :
                self.pushButton_25.setEnabled(True)
                QMessageBox.warning(self, 'لا توجد بيانات',  'لا توجد بيانات تخص المعلومات التي أدخلتها', QMessageBox.Ok)
                return
            
            self.tableWidget_3.setRowCount(0)
            self.tableWidget_3.insertRow(0)
            for row, form in enumerate(data):
                for col, item in enumerate(form):
                    self.tableWidget_3.setItem(row, col, QTableWidgetItem(str(item)))
                    col += 1
                row_pos = self.tableWidget_3.rowCount()
                self.tableWidget_3.insertRow(row_pos)    
        
        self.lineEdit_77.setText('')
        self.lineEdit_78.setText('')
        self.lineEdit_79.setText('') 

        self.pushButton_25.setEnabled(True)
        self.pushButton_18.setEnabled(False)
        self.pushButton_20.setEnabled(True)
        self.pushButton_21.setEnabled(True)        

    def egarat_edit(self):
        eg_id = self.lineEdit_77.text()
        eg_owner = self.lineEdit_73.text()
        eg_owner_id = self.lineEdit_68.text()
        eg_tenant = self.lineEdit_69.text()        
        eg_tenant_id = self.lineEdit_72.text()
        eg_date = self.dateEdit_10.date()
        eg_date = eg_date.toString(QtCore.Qt.ISODate)
        eg_category = self.comboBox_4.currentText()
        eg_discription = self.lineEdit_75.text()
        eg_address = self.lineEdit_76.text()
        eg_cont_term = self.lineEdit_70.text()
        eg_start = self.dateEdit_12.date()
        eg_start = eg_start.toString(QtCore.Qt.ISODate)
        eg_end = self.dateEdit_11.date()
        eg_end = eg_end.toString(QtCore.Qt.ISODate)
        eg_rental_value = self.lineEdit_74.text()
        eg_insurance = self.lineEdit_71.text()
        eg_purpose = self.lineEdit_67.text()

        self.cur.execute('''
        UPDATE rents SET owner=%s, owner_id=%s, tenant=%s, tenant_id=%s, date=%s, category=%s, discription=%s, address=%s, cont_term=%s, start=%s, end=%s, rental_value=%s, insurance=%s, purpose=%s
        WHERE id=%s''', (eg_owner, eg_owner_id, eg_tenant, eg_tenant_id, eg_date, eg_category, eg_discription , eg_address, eg_cont_term, eg_start, eg_end, eg_rental_value, eg_insurance, eg_purpose, eg_id))

        self.db.commit()

        self.lineEdit_77.setText('')
        self.egarat_search()        
        QMessageBox.warning(self, 'تعديل بيانات', 'تم تعديل البيانات بنجاح', QMessageBox.Ok)
        return


    def egarat_return(self):        
        self.egarat_search()
        self.pushButton_25.setEnabled(False)
        self.lineEdit_77.setEnabled(True)
        self.lineEdit_77.setText('')
        self.lineEdit_78.setText('')
        self.lineEdit_79.setText('')


    def contract_print(self):
        eg_owner = self.lineEdit_73.text()
        eg_owner_id = self.lineEdit_68.text()
        eg_tenant = self.lineEdit_69.text()        
        eg_tenant_id = self.lineEdit_72.text()
        eg_date = self.dateEdit_10.date()
        eg_date = eg_date.toString(QtCore.Qt.ISODate)        
        eg_category = self.comboBox_4.currentText()
        eg_discription = self.lineEdit_75.text()
        eg_address = self.lineEdit_76.text()
        eg_cont_term = self.lineEdit_70.text()
        eg_start = self.dateEdit_12.date()
        eg_start = eg_start.toString(QtCore.Qt.ISODate)
        eg_end = self.dateEdit_11.date()
        eg_end = eg_end.toString(QtCore.Qt.ISODate)
        eg_rental_value = self.lineEdit_74.text()
        eg_insurance = self.lineEdit_71.text()
        eg_purpose = self.lineEdit_67.text()        

        document = Document()
        style = document.styles['Normal']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(14)       

        title = arabic_reshaper.reshape('عقد إيجار أملاك')
        bidi_text = bidi.algorithm.get_display(title)        
        p = document.add_paragraph(bidi_text)
        
        date = arabic_reshaper.reshape(f'إنه في يوم {eg_date}')
        bidi_text = bidi.algorithm.get_display(date)        
        p = document.add_paragraph(bidi_text)

        first = arabic_reshaper.reshape(f'قد أجر السيد /   {eg_owner}')
        bidi_text = bidi.algorithm.get_display(first)        
        p = document.add_paragraph(bidi_text)

        first1 = arabic_reshaper.reshape(f'بطاقة رقم : {eg_owner_id}  مصري الجنسية')
        bidi_text = bidi.algorithm.get_display(first1)        
        p = document.add_paragraph(bidi_text)

        second = arabic_reshaper.reshape(f'إلى السيد /   {eg_tenant}')
        bidi_text = bidi.algorithm.get_display(second)        
        p = document.add_paragraph(bidi_text)

        second1 = arabic_reshaper.reshape(f'بطاقة رقم : {eg_tenant_id}  مصري الجنسية')
        bidi_text = bidi.algorithm.get_display(second1)        
        p = document.add_paragraph(bidi_text)

        type = arabic_reshaper.reshape(f'ما هو:   {eg_discription}')
        bidi_text = bidi.algorithm.get_display(type)        
        p = document.add_paragraph(bidi_text)

        for_use = arabic_reshaper.reshape(f'بقصد استخدامه:   {eg_purpose}')
        bidi_text = bidi.algorithm.get_display(for_use)        
        p = document.add_paragraph(bidi_text)

        address = arabic_reshaper.reshape(f' بالعقار رقم :   {eg_address}')
        bidi_text = bidi.algorithm.get_display(address)        
        p = document.add_paragraph(bidi_text)

        tamheed = arabic_reshaper.reshape(f'''وقد أقر المستأجر ( طرف ثان ) بأن العقار الجاري تأجيره مستوف لجميع الشروط اللازمة
                                        لاستخدامه لذات الغرض المتفق عليه.''')    
        bidi_text = bidi.algorithm.get_display(tamheed)        
        p = document.add_paragraph(bidi_text)

        tamheed_2 = arabic_reshaper.reshape(f'''وقد اتفق المتعاقدان وهما بكامل الأهلية على البنود الآتية.''')    
        bidi_text = bidi.algorithm.get_display(tamheed_2)        
        p = document.add_paragraph(bidi_text)

        band_1 = arabic_reshaper.reshape(f'''بند 1 : مدة  العقد هي سنتان فقط تبدأ من {eg_start}  وتنتهي في {eg_end} .''')    
        bidi_text = bidi.algorithm.get_display(band_1)        
        p = document.add_paragraph(bidi_text)

        band_2 = arabic_reshaper.reshape(f'''بند 2 : الأجرة المتفق عليها هي مبلغ {eg_rental_value} 
        يدفعها المستأجر مقدما ليد المالك أول كل شهر .''')    
        bidi_text = bidi.algorithm.get_display(band_2)        
        p = document.add_paragraph(bidi_text)

        band_3 = arabic_reshaper.reshape(f'''بند 3 : تقاضى الطرف الأول من الطرف الثاني مبلغ {eg_insurance} 
        كتأمين على محتويات العقار.''')
        bidi_text = bidi.algorithm.get_display(band_3)        
        p = document.add_paragraph(bidi_text)

        band_4 = arabic_reshaper.reshape (f'''بند 4 : إذا رغب أحد المتعاقدين إنهاء العلاقة الإيجارية في نهاية المدة المتعاقد عليها وجب أن   
        يعلن الطرف الآخر بخطاب مسجل مصحوب بعلم الوصول قبل إنتهاء مدة العقد بثلاثة 
        أشهر وإلا يعتبر العقد مجدد لمدة شهر فقط لحين التنبيه من أحد الطرفين.''')
        bidi_text = bidi.algorithm.get_display(band_4)        
        p = document.add_paragraph(bidi_text)

        band_5 = arabic_reshaper.reshape (f'''بند 5 : إذا تأخر المستأجر عن دفع الإيجار في المواعيد المحددة لمدة شهر فللمالك الحق أن 
        يلزمه بدفع الأجر والمصاريف وفسخ العقد بدون الحصول على حكم قضائي بعد التنبيه 
        عليه كتابة وقد قبل المستأجر بهذه الشروط.''')
        bidi_text = bidi.algorithm.get_display(band_5)        
        p = document.add_paragraph(bidi_text)

        band_6 = arabic_reshaper.reshape (f'''بند 6 : لا يجوز للمستأجر أن يؤجر العين المذكورة من الباطن أو يتنازل عنها للغير عن أي مدة 
        كانت أو إحداث تغيير للعين المؤجرة بدون إذن المالك كتابة وإذا خالف ذلك يحق للمالك  
        فسخ العقد وإلزام المستأجر بالعطل والأضرار والمصاريف التي تحدث.''')
        bidi_text = bidi.algorithm.get_display(band_6)        
        p = document.add_paragraph(bidi_text)

        band_7 = arabic_reshaper.reshape (f'''بند 7 : يلتزم المستأجر باستعمال العين المؤجرة له وفقا للبنود المتفق عليها بالعقد وأن يحافظ 
        عليها ويراعيها كما يراعي الإنسان ماله الخاص وإذا خالف ذلك يحق للمالك فسخ العقد 
        وإلزام المستأجر بالعطل والأضرار والمصاريف.''')
        bidi_text = bidi.algorithm.get_display(band_7)        
        p = document.add_paragraph(bidi_text)

        band_8 = arabic_reshaper.reshape (f'''بند 8 : جميع ما ينفقه المستأجر بعد استلام العين المؤجرة من دهانات أو لصق ورق أو ديكور 
        وخلافه لا يلزم المالك بشئ منها ولا يحق للمستأجر المطالبة بقيمتها عند الخروج منها بل 
        يكون متبرعا بها للمالك.''')
        bidi_text = bidi.algorithm.get_display(band_8)        
        p = document.add_paragraph(bidi_text)

        band_9 = arabic_reshaper.reshape (f'''بند 9 : يلتزم المستأجر بعمل الترميمات التأجيريه للعين المؤجرة من إصلاح للبلاظ أو الأبواب 
        والنوافذ والمفاتيح ودهان الحوائط وذلك طول مدة الإيجار أما الترميمات الضرورية تكون 
        على نفقة المالك.''')
        bidi_text = bidi.algorithm.get_display(band_9)        
        p = document.add_paragraph(bidi_text)

        band_10 = arabic_reshaper.reshape (f'''بند 10 : جميع ما يملكه الساكن من أثاث ومنقولات وبضائع وخلافه بالعين المؤجرة يحق للمالك 
        الحجز عليها في حالة التأخير عن دفع الإيجار لاستيفاء حقه منها.''')
        bidi_text = bidi.algorithm.get_display(band_10)        
        p = document.add_paragraph(bidi_text)

        band_11 = arabic_reshaper.reshape (f'''بند 11 : إذا ترك الساكن العين المؤجرة فيلزم دفع باقي المدة التعاقد عليها مع دفع قيمة ما يكون 
        قد أتلف منها.''')
        bidi_text = bidi.algorithm.get_display(band_11)        
        p = document.add_paragraph(bidi_text)

        band_12 = arabic_reshaper.reshape (f'''بند 12 : يلتزم المستأجر برد العين المؤجرة في حالة إنتهاء مدة التعاقد بالحالة التي كانت عليها 
        وقت التسليم ويتحمل كافة النفقات إذا حدث للعين تلف أو هلاك يرجع إلى خطأ المستأجر ''')
        bidi_text = bidi.algorithm.get_display(band_12)        
        p = document.add_paragraph(bidi_text)

        band_13 = arabic_reshaper.reshape (f'''.بند 13 : إذا حدث أمر مخل بالعين المؤجرة يحق للمالك إخراج المستأجر من العين بمجرد 
        التنبيه عليه شفويا ويحق له فسخ العقد.''')
        bidi_text = bidi.algorithm.get_display(band_13)        
        p = document.add_paragraph(bidi_text)

        band_14 = arabic_reshaper.reshape (f'''بند 14 : يلتزم المستأجر بدفع قيمة فواتير الكهرباء والغاز أما فواتير الماء يلتزم المالك بدفع 
        قيمتها.''')
        bidi_text = bidi.algorithm.get_display(band_14)        
        p = document.add_paragraph(bidi_text)

        band_15 = arabic_reshaper.reshape (f'''بند 15 : يخضع هذا العقد لأحكام القانون رقم 4 لسنة 1996 بشأن سريان الحكم القانوني المدني 
        على الأماكن التي لم يسبق تأجيرها والأماكن التي انتهت أو تنتهي عقود إيجارها.''')
        bidi_text = bidi.algorithm.get_display(band_15)        
        p = document.add_paragraph(bidi_text)

        band_16 = arabic_reshaper.reshape (f'''بند 16 : يتم استرداد مبلغ التأمين في نهاية المدة الإيجارية بعد معاينة الشقة وخصم قيمة التلفيات ''')
        bidi_text = bidi.algorithm.get_display(band_16)        
        p = document.add_paragraph(bidi_text)

        band_17 = arabic_reshaper.reshape (f'''بند 17 : تختص محكمة المنصورة الابتدائية بالنظر فيما قد ينشأ من منازعات أو خلافه حول    
        بنود العقد. وقد تحرر هذا العقد من نسختين بيد كل من الطرفين نسخة للعمل بموجبها.''')
        bidi_text = bidi.algorithm.get_display(band_17)        
        p = document.add_paragraph(bidi_text)

        end = arabic_reshaper.reshape (f'''     الطرف الأول ( المالك )                                  الطرف الثاني ( المستأجر ) ''')
        bidi_text = bidi.algorithm.get_display(end)        
        p = document.add_paragraph(bidi_text)

        signe = arabic_reshaper.reshape (f'''{eg_owner}                              {eg_tenant}  ''')
        bidi_text = bidi.algorithm.get_display(signe)        
        p = document.add_paragraph(bidi_text)
        
        #p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT        
        document.save('egarat.docx')
        QMessageBox.warning(self, 'نسخ بيانات', 'تم نسخ البيانات في ملف: egarat.docx بنجاح', QMessageBox.Ok)
        return

        # with open('egarat.txt', 'w') as file:
        #     file.write(multi_line_string)

        #file_name = 'egarat.txt'
        #win32api.ShellExecute(0, "print", 'egarat.docx', None, ".", 0)
    def tamleek_tap(self):
        self.pushButton.setStyleSheet("background-color : ")
        self.pushButton_2.setStyleSheet("background-color : ")
        self.pushButton_3.setStyleSheet("background-color : ")
        self.pushButton_4.setStyleSheet("background-color : #c2fc03")
        self.tabWidget.setCurrentIndex(3)

    def tamleek_save(self):
        ta_seller_id = self.lineEdit_30.text()
        ta_seller = self.lineEdit_28.text()
        ta_buyer = self.lineEdit_26.text()
        ta_buyer_id = self.lineEdit_27.text()
        ta_date = self.dateEdit_6.date()
        ta_date = ta_date.toString(QtCore.Qt.ISODate)
        ta_category = self.comboBox_2.currentText()
        ta_discription = self.lineEdit_31.text()
        ta_address = self.lineEdit_29.text()
        ta_title_deep = self.lineEdit_33.text() # سند الملكية
        ta_price = self.lineEdit_32.text()
        ta_advance = self.lineEdit_48.text()
        ta_remain = self.lineEdit_49.text()
        ta_installments = self.lineEdit_35.text()
        ta_install_value = self.lineEdit_36.text()
        ta_penalty = self.lineEdit_34.text() # الشرط الجزائي

        if ta_seller =='' or ta_seller_id =='' \
        or ta_buyer =='' or ta_buyer_id =='' \
        or ta_category =='' or ta_discription =='' \
        or ta_address == '' or ta_title_deep=='' \
        or ta_price == '' or ta_advance == '' \
        or ta_remain == '' or ta_installments == '' \
        or ta_install_value == '' or ta_penalty == '' :        

            QMessageBox.warning(self, 'بيانات ناقصة', 'من فضلك تأكد من إدخال جميع البيانات', QMessageBox.Ok)
            return

        self.cur.execute('''
            INSERT INTO ownership(seller, seller_id, buyer, buyer_id, date, category, discription, address, title_deed, price, advance, remain, installments, install_value, penalty)
            VALUES(%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
            ''',(ta_seller, ta_seller_id, ta_buyer, ta_buyer_id, ta_date, ta_category, ta_discription, ta_address, ta_title_deep, ta_price, ta_advance, ta_remain, ta_installments, ta_install_value, ta_penalty))

        self.db.commit()
        self.tamleek_search()
        self.pushButton_8.setEnabled(False)
        self.pushButton_15.setEnabled(True)
        QMessageBox.warning(self, 'حفظ بيانات', 'تم حفظ البيانات بنجاح', QMessageBox.Ok)
        return

    def tamleek_clrscreen(self):

        self.lineEdit_28.setText('')
        self.lineEdit_30.setText('')
        self.lineEdit_26.setText('')
        self.lineEdit_27.setText('')
        self.lineEdit_31.setText('')
        self.lineEdit_29.setText('')
        self.lineEdit_32.setText('')
        self.lineEdit_33.setText('')
        self.lineEdit_34.setText('')
        self.lineEdit_35.setText('')
        self.lineEdit_36.setText('')
        self.lineEdit_48.setText('')
        self.lineEdit_49.setText('')
        self.pushButton_8.setEnabled(True)
        self.pushButton_15.setEnabled(False)
        self.pushButton_16.setEnabled(False)
        self.pushButton_17.setEnabled(False)

    def tamleek_search(self):

            self.tableWidget_4.setRowCount(0)
            self.tableWidget_4.insertRow(0)
            self.cur.execute('''
            SELECT id, seller, buyer, date, category, price, advance, remain, installments, install_value FROM ownership
            ''')
            data = self.cur.fetchall()

            for row, form in enumerate(data):
                for col, item in enumerate(form):
                    self.tableWidget_4.setItem(row, col, QTableWidgetItem(str(item)))
                    col += 1
                row_pos = self.tableWidget_4.rowCount()
                self.tableWidget_4.insertRow(row_pos)

            self.pushButton_26.setEnabled(False)

    def tamleek_filter(self):
        
        if self.lineEdit_42.text() == '' and self.lineEdit_41.text()=='' and self.lineEdit_47.text()=='':
            QMessageBox.warning(self, 'لا توجد بيانات',  'لا توجد بيانات يمكن البحث عنها', QMessageBox.Ok)
            return
        if self.lineEdit_42.text() != '':            
            ta_id = self.lineEdit_42.text()
            sql =f''' SELECT * FROM ownership WHERE id='{ta_id}' '''
                        
            self.cur.execute(sql)
            data = self.cur.fetchone()
            if data == None :
                QMessageBox.warning(self, 'لا توجد بيانات',  'لا توجد بيانات تخص المعلومات التي أدخلتها', QMessageBox.Ok)
                return

            self.tabWidget_5.setCurrentIndex(0)        
            self.lineEdit_28.setText(data[1])
            self.lineEdit_30.setText(str(data[2]))       
            self.lineEdit_26.setText(data[3]) 
            self.lineEdit_27.setText(str(data[4]))
            self.dateEdit_6.setDate(data[5]) 
            self.comboBox_2.setCurrentText(data[6])
            self.lineEdit_31.setText(data[7])
            self.lineEdit_29.setText(data[8])        
            self.lineEdit_33.setText(data[9])
            self.lineEdit_32.setText(data[10])
            self.lineEdit_48.setText(data[11])
            self.lineEdit_49.setText(data[12])
            self.lineEdit_35.setText(data[13])
            self.lineEdit_36.setText(data[14])
            self.lineEdit_34.setText(data[15])

        else:            
            if self.lineEdit_41.text() != '':
                ta_name = self.lineEdit_41.text()
                sql = f''' SELECT * FROM ownership WHERE seller LIKE '%{ta_name}%' '''
            elif self.lineEdit_47.text() != '':
                category = self.lineEdit_47.text()
                sql = f''' SELECT * FROM ownership WHERE category LIKE '%{category}%' '''
            self.cur.execute(sql)
            data = self.cur.fetchall()            
            if data == [] :
                self.pushButton_23.setEnabled(True)
                QMessageBox.warning(self, 'لا توجد بيانات',  'لا توجد بيانات تخص المعلومات التي أدخلتها', QMessageBox.Ok)
                return
            
            self.tableWidget_4.setRowCount(0)
            self.tableWidget_4.insertRow(0)
            for row, form in enumerate(data):
                for col, item in enumerate(form):
                    self.tableWidget_4.setItem(row, col, QTableWidgetItem(str(item)))
                    col += 1
                row_pos = self.tableWidget_4.rowCount()
                self.tableWidget_4.insertRow(row_pos)
        
        self.lineEdit_41.setText('')
        self.lineEdit_42.setText('')
        self.lineEdit_47.setText('')

        self.pushButton_17.setEnabled(True)           
        self.pushButton_26.setEnabled(True)
        self.pushButton_8.setEnabled(False)
        self.pushButton_15.setEnabled(True)
    
    def tamleek_edit(self):

        ta_id = self.lineEdit_42.text()
        ta_seller = self.lineEdit_28.text()
        ta_seller_id = self.lineEdit_30.text()
        ta_buyer = self.lineEdit_26.text()        
        ta_buyer_id = self.lineEdit_27.text()
        ta_date = self.dateEdit_6.date()
        ta_date = ta_date.toString(QtCore.Qt.ISODate)
        ta_category = self.comboBox_2.currentText()
        ta_discription = self.lineEdit_31.text()
        ta_address = self.lineEdit_29.text()
        ta_title_deed = self.lineEdit_33.text()
        ta_price = self.lineEdit_32.text()
        ta_advance = self.lineEdit_48.text()        
        ta_remain = self.lineEdit_49.text()
        ta_installments = self.lineEdit_35.text()
        ta_install_value = self.lineEdit_36.text()
        ta_penalty = self.lineEdit_34.text()

        sql = f''' UPDATE ownership SET\
        seller='{ta_seller}', seller_id='{ta_seller_id}',\
        buyer='{ta_buyer}', buyer_id='{ta_buyer_id}',\
        date='{ta_date}', category='{ta_category}',\
        discription='{ta_discription}', address='{ta_address}',\
        title_deed='{ta_title_deed}', price='{ta_price}',\
        advance='{ta_advance}', remain='{ta_remain}',\
        installments='{ta_installments}', install_value='{ta_install_value}',\
        penalty='{ta_penalty}' WHERE id='{ta_id}' '''
        self.cur.execute(sql)
        self.db.commit()

        self.lineEdit_42.setText('')
        self.tamleek_search()        
        QMessageBox.warning(self, 'تعديل بيانات', 'تم تعديل البيانات بنجاح', QMessageBox.Ok)
        return
    
    def tamleek_return(self):        
        self.tamleek_search()
        self.pushButton_26.setEnabled(False)
        self.lineEdit_42.setEnabled(True)
        self.lineEdit_41.setText('')
        self.lineEdit_42.setText('')
        self.lineEdit_47.setText('')

    def tamleek_print(self):
        ta_seller = self.lineEdit_28.text()
        ta_seller_id = self.lineEdit_30.text()
        ta_buyer = self.lineEdit_26.text()
        ta_buyer_id = self.lineEdit_27.text()
        ta_date = self.dateEdit_2.date()
        ta_date = ta_date.toString(QtCore.Qt.ISODate)
        #ta_date = ta_date.toString('dd/MM/yyyy')        
        ta_category = self.comboBox_2.currentText()
        ta_discription = self.lineEdit_31.text()
        ta_address = self.lineEdit_29.text()
        ta_title_deed = self.lineEdit_33.text()
        ta_price = self.lineEdit_32.text()
        ta_advance = self.lineEdit_33.text()
        ta_installments = self.lineEdit_35.text()
        ta_remain = self.lineEdit_34.text()
        ta_install_value = self.lineEdit_36.text()
        ta_penalty = self.lineEdit_34.text()

        document = Document()
        style = document.styles['Normal']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(14)
        #font.rtl = True
        #document.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

        title = arabic_reshaper.reshape('عقد بيع عقار')
        bidi_text = bidi.algorithm.get_display(title)        
        p = document.add_paragraph(bidi_text)

        date = arabic_reshaper.reshape(f'إنه في يوم {ta_date} قد تحرر هذا العقد بين كل من : -')
        bidi_text = bidi.algorithm.get_display(date)        
        p = document.add_paragraph(bidi_text)

        first = arabic_reshaper.reshape(f'''السيد /  {ta_seller}   مصري الجنسية 
بطاقة رقم  /  {ta_seller_id}   طرف أول /  بائع ''')
        bidi_text = bidi.algorithm.get_display(first)        
        p = document.add_paragraph(bidi_text)

        second = arabic_reshaper.reshape(f'''والسيد /  {ta_buyer}    مصري الجنسية 
بطاقة رقم  /  {ta_buyer_id}   طرف ثان /  مشتري ''')
        bidi_text = bidi.algorithm.get_display(second)        
        p = document.add_paragraph(bidi_text)        

        tamheed = arabic_reshaper.reshape(f'''وبعد أن أقر الطرفان بأهليتهما للتعاقد والتصرف والإبرام فقد اتفقا علي الآتي :''')
        bidi_text = bidi.algorithm.get_display(tamheed)        
        p = document.add_paragraph(bidi_text)

        band_1 = arabic_reshaper.reshape(f'''بند 1: اتفق الطرف الأول مع الطرف الثاني علي أن يبيع الطرف الأول للطرف الثاني 
        العقار :  {ta_discription} 
        بالعنوان :  {ta_address}''')
        bidi_text = bidi.algorithm.get_display(band_1)        
        p = document.add_paragraph(bidi_text)

        band_2 = arabic_reshaper.reshape(f'''بند 2: باع وأسقط وتنازل الطرف الأول وبكافة الضمانات الفعلية والقانونية والمثبتة 
        للملكية إلي الطرف الثاني العقار (محل القعد) بالعنوان والمواصفات أعلاه. ''')
        bidi_text = bidi.algorithm.get_display(band_2)        
        p = document.add_paragraph(bidi_text)

        band_3 = arabic_reshaper.reshape(f'''بند 3: يقر الطرف الأول البائع بأن ملكية العقار المباع قد آلت إليه عن طريق :
        سند الملكية : {ta_title_deed}''')
        bidi_text = bidi.algorithm.get_display(band_3)        
        p = document.add_paragraph(bidi_text)

        band_4 = arabic_reshaper.reshape(f'''بند 4: تم بيع العقار (محل العقد) نظير
        ثمن إجمالي : {ta_price}
        وقد قام الطرف الثاني المشترى وقت تحرير هذا العقد بسداد كامل ثمن البيع يعتبر  
        توقيع الطرف الأول البائع على هذا العقد بمثابة إقرار منه باستلام الثمن بالكامل.''')
        bidi_text = bidi.algorithm.get_display(band_4)        
        p = document.add_paragraph(bidi_text)

        band_5 = arabic_reshaper.reshape(f'''بند 5: قام الطرف الأول البائع بتسليم العقار المباع للطرف الثاني المشترى والذي قام 
        بمعاينته المعاينة التامة الكاملة النافية للجهالة شرعاً وقانوناً وقبل شراءه بالحالة 
        التي عليها ويعتبر توقيع المشترى على هذا العقد بمثابة إقرار منه بالاستلام . ''')
        bidi_text = bidi.algorithm.get_display(band_5)        
        p = document.add_paragraph(bidi_text)

        band_6 = arabic_reshaper.reshape(f'''بند 6: يقر الطرف الأول البائع بأن العقار المباع ليس عليه أية مستحقات للغير أو لأي 
        جهة حكومية ومنها الضرائب العقارية والكهرباء والمياه وخلافه وإذا ظهر خلاف 
        ذلك يتحمل البائع وحده قيمة المطالبات وذلك حتى تاريخ تحرير هذا العقد .''')
        bidi_text = bidi.algorithm.get_display(band_6)        
        p = document.add_paragraph(bidi_text)

        band_7 = arabic_reshaper.reshape(f'''بند 7: يقر الطرف الأول البائع بأن العقار المباع خالي من كافة الحقوق العينية أياً كان 
        نوعها كرهن الاختصاص أو الامتياز وحقوق الانتفاع والارتفاق ظاهرة أو خفية .''')
        bidi_text = bidi.algorithm.get_display(band_7)        
        p = document.add_paragraph(bidi_text)

        band_8 = arabic_reshaper.reshape(f'''بند 8: في حالة إخلال أي طرف بأي من التزاماته المنصوص عليها في هذا العقد يلتزم 
        الطرف المخل بدفع شرطا جزائيا قدره : {ta_penalty}   
        تعويضاً اتفاقياً لا رقابة للقضاء عليه.''')
        bidi_text = bidi.algorithm.get_display(band_8)        
        p = document.add_paragraph(bidi_text)

        band_9 = arabic_reshaper.reshape(f'''بند 9: تختص محكمة المنصورة الابتدائية بالفصل في أي نزاع قد ينشأ لا قدر الله بين 
        الطرفين بسبب هذا العقد .''')
        bidi_text = bidi.algorithm.get_display(band_9)        
        p = document.add_paragraph(bidi_text)

        band_10 = arabic_reshaper.reshape(f'''بند 10: جميع مصروفات هذا العقد وأتعاب تحرير العقد النهائي ورسوم تسجيله وكافة ما 
            يلزم أو يقتضى الأمر إنفاقه في هذا الشأن تكون على عاتق الطرف الثاني .''')
        bidi_text = bidi.algorithm.get_display(band_10)        
        p = document.add_paragraph(bidi_text)

        band_11 = arabic_reshaper.reshape(f'''بند 11: تحرر هذا العقد من نسختين بيد كل طرف نسخة للعمل بموجبها عند اللزوم . ''')
        bidi_text = bidi.algorithm.get_display(band_11)        
        p = document.add_paragraph(bidi_text)

        band_12 = arabic_reshaper.reshape(f'''                   طرف ثان :	                                  طرف أول :  ''')
        bidi_text = bidi.algorithm.get_display(band_12)        
        p = document.add_paragraph(bidi_text)

        band_13 = arabic_reshaper.reshape(f'''{ta_seller}                                    {ta_buyer}''')
        bidi_text = bidi.algorithm.get_display(band_13)        
        p = document.add_paragraph(bidi_text)        
        document.save('tamleek.docx')
        QMessageBox.warning(self, 'نسخ بيانات', 'تم نسخ البيانات في ملف: tamleek.docx بنجاح', QMessageBox.Ok)
        return

        #win32api.ShellExecute(0, "print", 'tamleek.docx', None, ".", 0)


    def login(self):
        username = self.lineEdit_43.text()
        pass_word = self.lineEdit_44.text()

        sql = f''' SELECT * FROM login WHERE user_name = '{username}' and password = '{pass_word}' '''
        
        self.cur.execute(sql)
        data = self.cur.fetchone()        
        if data == None :            
            QMessageBox.warning(self, 'بيانات خاطئة',  'معلومات غير صحيحة ', QMessageBox.Ok)
            self.lineEdit_43.setText('')
            self.lineEdit_44.setText('')
            return
        if data[3] == 1:
            self.pushButton_29.setEnabled(True)
            self.checkBox.setVisible(True)
        else:
            self.pushButton_29.setEnabled(False)
            self.checkBox.setVisible(False)

        self.groupBox.setEnabled(True)
        self.lineEdit_43.setText('')
        self.lineEdit_44.setText('')
        self.pushButton_27.setEnabled(False)
        

    def signup(self):
        user = self.lineEdit_43.text()
        pass_word = self.lineEdit_44.text()
        if user == '' or pass_word == '' :
            QMessageBox.warning(self, 'بيانات خاطئة',  'معلومات غير صحيحة ', QMessageBox.Ok)
            self.lineEdit_43.setText('')
            self.lineEdit_44.setText('')
            return

        sql = f''' SELECT * FROM login WHERE user_name = '{user}' '''
        
        self.cur.execute(sql)
        data = self.cur.fetchone()        
        if data != None :
            QMessageBox.warning(self, 'بيانات مكررة',  'هذا المستخدم موجود بالفعل في قاعدة البيانات ', QMessageBox.Ok)
            # self.lineEdit_43.setText('')
            # self.lineEdit_44.setText('')
            return

        if self.checkBox.isChecked():
            admins = 1
        else:
            admins = 0

        self.cur.execute('''
        INSERT INTO login (user_name, password, admin)
        VALUES(%s, %s, %s) ''',(user, pass_word, admins))

        self.db.commit()

        self.lineEdit_43.setText('')
        self.lineEdit_44.setText('')
        self.pushButton_29.setEnabled(False)
        self.checkBox.setVisible(False)



    def signout(self):
        self.tabWidget.setCurrentIndex(4)
        self.tabWidget.tabBar().setVisible(False)
        self.groupBox.setEnabled(False)
        self.pushButton_27.setEnabled(True)
        self.pushButton_29.setEnabled(False)
        self.checkBox.setVisible(False)

def main():
    app = QApplication(sys.argv)
    Window = Main()
    Window.show()
    app.exec_()
if __name__ == '__main__':
    main()